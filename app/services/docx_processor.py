import os
import json
import re
import hashlib
from typing import Optional
from app.services.aes_gcm import encrypt_with_metadata, generate_key as generate_key_aes
from docx import Document
from app.services.pii_main import extract_all_pii

def mask_docx_sensitive_text(docx_path: str, key_path: Optional[str] = None, enabled_pii_categories=None):
    if key_path is None:
        key_path = docx_path.replace(".docx", ".key")

    if os.path.exists(key_path):
        with open(key_path, "rb") as f:
            key = f.read()
    else:
        key = generate_key_aes()
        with open(key_path, "wb") as f:
            f.write(key)

    document = Document(docx_path)

    full_text = ""
    for para in document.paragraphs:
        full_text += para.text + "\n"

    all_pii_list = extract_all_pii(full_text, enabled_pii_categories)

    # Filter by enabled categories
    non_selectable_categories = ['IC', 'Email', 'DOB', 'Bank Account', 'Passport', 'Phone', 'Credit Card', 'Address', 'Vehicle Registration']
    
    # Map Ollama category names to expected format
    category_mapping = {
        'ACCOUNT': 'Bank Account',
        'EMAIL': 'Email',
        'PHONE': 'Phone',
        'CREDIT_CARD': 'Credit Card',
        'NAMES': 'NAMES',
        'ORG_NAMES': 'ORG_NAMES',
        'ETHNIC': 'ETHNIC',
        'DOB': 'DOB',
        'PASSPORT': 'Passport',
        'ADDRESS': 'Address',
        'VEHICLE_REGISTRATION': 'Vehicle Registration',
        # Legacy category mappings
        'RACES': 'ETHNIC',
        'RELIGIONS': 'ETHNIC',
        'LOCATIONS': 'Address',
        'STATUS': 'NAMES'
    }
    
    filtered_pii_list = []
    for label, value in all_pii_list:
        mapped_label = category_mapping.get(label, label)
        if mapped_label in enabled_pii_categories or label in non_selectable_categories or mapped_label in non_selectable_categories:
            filtered_pii_list.append((label, value))
            print(f"[FILTER] Masking DOCX: {label} = {value}")
        else:
            print(f"[FILTER] Skipping DOCX (not enabled): {label} (mapped: {mapped_label})")

    all_pii_list = filtered_pii_list

    unique_pii = {}
    masked_pii = []

    for label, value in all_pii_list:
        if value not in unique_pii:
            try:
                encrypted = encrypt_with_metadata(value, key)
                encrypted_str = json.dumps(encrypted)
                value_hash = hashlib.md5(value.encode()).hexdigest()[:8]
                unique_tag = f"[ENC:{label}_{value_hash}]"

                unique_pii[value] = {"encrypted": encrypted_str, "tag": unique_tag}
                masked_pii.append({
                    "original": value,
                    "encrypted": encrypted_str,
                    "label": label,
                    "masked": unique_tag
                })
            except Exception as e:
                print(f"[ERROR] Failed to encrypt '{value}': {e}")

    sorted_pii_items = sorted(unique_pii.items(), key=lambda x: len(x[0]), reverse=True)

    for para in document.paragraphs:
        if para.text.strip():
            masked_text = para.text

            for pii_value, pii_info in sorted_pii_items:
                # Use word boundary regex to replace only complete words
                escaped_pii = re.escape(pii_value)
                pattern = r'\b' + escaped_pii + r'\b'
                masked_text = re.sub(pattern, pii_info["tag"], masked_text)

            para.text = masked_text

    masked_path = docx_path.replace(".docx", ".masked.docx")
    document.save(masked_path)

    json_path = docx_path.replace(".docx", ".masked.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(masked_pii, f, ensure_ascii=False, indent=2)

    return masked_path, json_path, key_path

def run_docx_processing(docx_path: str, enabled_pii_categories=None):
    try:
        key_path = docx_path.replace(".docx", ".key")
        masked_docx, json_path, key_file = mask_docx_sensitive_text(
            docx_path, key_path, enabled_pii_categories
        )
        return {
            "status": "success",
            "masked_docx": masked_docx,
            "json_output": json_path,
            "key_file": key_file
        }
    except Exception as e:
        return {
            "status": "error",
            "message": str(e)
        }
